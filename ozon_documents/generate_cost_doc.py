from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面设置 ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.2)

# ── 样式辅助 ──────────────────────────────────────────────────────────────────
BLUE_DARK  = (0x1F, 0x49, 0x7D)
BLUE_MID   = (0x2E, 0x74, 0xB5)
BLUE_LIGHT = (0xBD, 0xD7, 0xEE)
GRAY_BG    = (0xF2, 0xF2, 0xF2)
CODE_GRAY  = (0x59, 0x59, 0x59)

def rgb(r, g, b):
    return RGBColor(r, g, b)

def set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    if isinstance(rgb, str):
        hex_color = rgb
    else:
        hex_color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, border_color='AAAAAA'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), border_color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(20)
    run.font.color.rgb = rgb(*BLUE_DARK)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = rgb(*BLUE_MID)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = rgb(*BLUE_DARK)
    return p

def body(text, bold_parts=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(16)
    if bold_parts:
        # simple: split on ** markers
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
            run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5)
    # parse bold
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)

def blockquote(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        # left border via paragraph border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '12')
        left.set(qn('w:space'), '6')
        left.set(qn('w:color'), '2E74B5')
        pBdr.append(left)
        pPr.append(pBdr)
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F0F0')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = rgb(*CODE_GRAY)

def add_table(headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # set col widths
    usable = section.page_width - section.left_margin - section.right_margin
    if col_widths:
        widths = [int(usable * w) for w in col_widths]
    else:
        w = usable // n_cols
        widths = [w] * n_cols

    # header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr_cells[i]
        c.width = widths[i]
        set_cell_bg(c, BLUE_LIGHT)
        set_cell_border(c, '2E74B5')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = rgb(*BLUE_DARK)

    # data rows
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg = GRAY_BG if ri % 2 == 1 else None
        for ci, val in enumerate(row):
            c = cells[ci]
            c.width = widths[ci]
            if bg:
                set_cell_bg(c, bg)
            set_cell_border(c, 'CCCCCC')
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            # first col: left align; others: left
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            is_bold = val.startswith('**') and val.endswith('**')
            clean = val.strip('*')
            run = p.add_run(clean)
            run.font.size = Pt(9.5)
            if is_bold:
                run.bold = True
            # code-like: monospace for field names (backtick style)
            if val.startswith('`') or (ci == 0 and '_' in val and ' ' not in val.strip('`')):
                run.font.name = 'Courier New'
                run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table

# ══════════════════════════════════════════════════════════════════════════════
# 文档内容
# ══════════════════════════════════════════════════════════════════════════════

h1('SKU 成本表结构设计')

p = doc.add_paragraph()
run = p.add_run('电商管理平台 · 数据库设计方案')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p.paragraph_format.space_after = Pt(12)

# ── 背景 ──────────────────────────────────────────────────────────────────────
h2('一、背景')
body('基于现有「成本模板」Google Sheet，将 SKU 成本数据结构化落库，支持按平台、配送模式、公司规模四个维度计算每个 SKU 的成本与利润。')

# ── 设计思路 ──────────────────────────────────────────────────────────────────
h2('二、设计思路')
body('成本字段分两类，对应两张表：')
bullet('**固定成本**：采购价、运费到白俄、清关、关税、进口增值税、包装尺寸。与平台无关，每个 SKU 只有一份。')
bullet('**可变参数**：平台运费、佣金率、退货率、存储费、广告费率、税率、汇率、目标售价。随平台（WB / OZON）、配送模式（FBO / FBS）、公司规模（一般纳税人 / 小规模）而不同。')

doc.add_paragraph()

# ── 表一 ──────────────────────────────────────────────────────────────────────
h2('三、表一：ec_sku_costs（每个 SKU 一条）')
body('存采购链路中不随平台变化的固定成本。')

add_table(
    ['字段', '类型', '说明'],
    [
        ['offer_id',           'string, 唯一', 'SKU 编码，如 KJ-228-BK'],
        ['product_name',       'string',       '商品名称'],
        ['purchase_price_cny', 'decimal',      '国内采货价 / 件（含税到火车站）'],
        ['freight_to_by_cny',  'decimal',      '到白俄运费 / 件（按外箱体积换算）'],
        ['customs_misc_cny',   'decimal',      '清关杂费 / 件'],
        ['customs_duty_cny',   'decimal',      '关税 / 件'],
        ['import_vat_cny',     'decimal',      '进口增值税 / 件（通常为前四项合计 × 20%）'],
        ['goods_cost_cny',     'decimal',      '货物成本 / 件（以上五项合计，冗余存储便于直接读取）'],
        ['pkg_length_cm',      'decimal',      '外包装长（cm）'],
        ['pkg_width_cm',       'decimal',      '外包装宽（cm）'],
        ['pkg_height_cm',      'decimal',      '外包装高（cm）'],
        ['pkg_volume_l',       'decimal',      '包装容量（L），用于计算平台运费'],
        ['damage_rate',        'decimal',      '货损率，如 0.01'],
        ['misc_cost_cny',      'decimal',      '杂费 / 件'],
        ['memo',               'text',         '备注'],
    ],
    col_widths=[0.32, 0.18, 0.50]
)

# ── 表二 ──────────────────────────────────────────────────────────────────────
h2('四、表二：ec_sku_platform_costs（每个 SKU × 场景一条）')
body('唯一索引：(offer_id, platform, delivery_mode, company_type)')
body('每个 SKU 最多有 8 个场景（2 平台 × 2 配送模式 × 2 公司类型），实际常用 4～6 个。')

h3('4.1 维度字段')
add_table(
    ['字段', '类型', '枚举值', '说明'],
    [
        ['offer_id',      'string', '—',                    '关联 ec_sku_costs'],
        ['platform',      'string', 'wb / ozon',            '平台'],
        ['delivery_mode', 'string', 'fbo / fbs',            '配送模式'],
        ['company_type',  'string', 'general / small',      '一般纳税人 / 小规模'],
    ],
    col_widths=[0.27, 0.15, 0.22, 0.36]
)

h3('4.2 平台运费参数')
add_table(
    ['字段', '类型', '说明'],
    [
        ['base_logistics_rub', 'decimal', '基础运费（RUB），WB 按包装体积计算得出'],
        ['logistics_coeff',    'decimal', '仓库运费系数，如 1.55'],
        ['fbo_delivery_cny',   'decimal', 'FBO 送仓费（CNY）；FBS 填 0'],
    ],
    col_widths=[0.32, 0.18, 0.50]
)
blockquote(['平台运费 = base_logistics_rub × logistics_coeff ÷ 汇率 + fbo_delivery_cny'])

h3('4.3 退货参数')
add_table(
    ['字段', '类型', '说明'],
    [
        ['return_rate',         'decimal', '退货率，如 0.10（即 10% 退货）'],
        ['return_trip_cny',     'decimal', '退货返程费 / 次（CNY）'],
        ['return_freight_cny',  'decimal', '退货运费 / 次（CNY）'],
    ],
    col_widths=[0.32, 0.18, 0.50]
)
blockquote([
    '退货分摊 / 件 = （平台运费 + 返程费 + 退货运费）× 退货率 ÷（1 − 退货率）',
    '逻辑：发 10 件，9 件成交，1 件退货；9 件成交共同分摊 1 次退货产生的全部费用。',
])

h3('4.4 其他平台费用')
add_table(
    ['字段', '类型', '说明'],
    [
        ['storage_30d_cny',  'decimal', '30 天存储费（CNY）；FBS 填 0'],
        ['acquiring_rate',   'decimal', '收单费率，如 0.017'],
        ['commission_rate',  'decimal', '平台佣金率（按商品类目不同，WB 约 21%，OZON 约 7.5%）'],
        ['ad_spend_rate',    'decimal', '广告费率（占售价比例），建议控制在 10% 以内'],
    ],
    col_widths=[0.32, 0.18, 0.50]
)

h3('4.5 税务')
add_table(
    ['字段', '类型', '说明'],
    [
        ['sales_tax_rate',        'decimal', '销售税率：一般纳税人 0.20（增值税），小规模 0.06（营业税）'],
        ['import_vat_deductible', 'boolean', '进口增值税是否可抵扣：一般纳税人 true，小规模 false'],
    ],
    col_widths=[0.32, 0.18, 0.50]
)

h3('4.6 汇率与目标价')
add_table(
    ['字段', '类型', '说明'],
    [
        ['exchange_rate_rub_cny', 'decimal', '卢布 / 人民币汇率，定期手动更新'],
        ['target_price_rub',      'decimal', '目标售价（RUB）'],
        ['min_price_rub',         'decimal', '最低可接受售价（RUB）'],
    ],
    col_widths=[0.32, 0.18, 0.50]
)

# ── 计算公式 ──────────────────────────────────────────────────────────────────
h2('五、计算公式（运行时计算，不落库）')
body('以下字段由代码根据以上参数实时计算，不存入数据库：')

code_block([
    'revenue_cny          = target_price_rub ÷ exchange_rate',
    '',
    'platform_freight_cny = base_logistics_rub × logistics_coeff ÷ exchange_rate + fbo_delivery_cny',
    'return_cost_cny      = (platform_freight + return_trip + return_freight)',
    '                       × return_rate ÷ (1 − return_rate)',
    'acquiring_cny        = revenue × acquiring_rate',
    'commission_cny       = revenue × commission_rate',
    'ad_spend_cny         = revenue × ad_spend_rate',
    'sales_tax_cny        = revenue × sales_tax_rate',
    'damage_cost_cny      = goods_cost × damage_rate',
    'import_vat_offset    = import_vat_deductible ? import_vat_cny : 0',
    '',
    'total_cost_cny = goods_cost + platform_freight + return_cost + storage_30d',
    '               + acquiring + commission + ad_spend + sales_tax',
    '               + damage_cost + misc_cost − import_vat_offset',
    '',
    'profit_cny  = revenue_cny − total_cost_cny',
    'margin      = profit_cny ÷ revenue_cny',
])

doc.add_paragraph()

# ── 场景示例 ──────────────────────────────────────────────────────────────────
h2('六、场景示例')
body('以 KJ-228-SV（毛巾架）为例，OZON / FBO / 一般纳税人场景：')

add_table(
    ['项目', '金额（CNY）'],
    [
        ['货物成本',           '454.1'],
        ['平台运费（FBO）',    '45.0'],
        ['退货分摊',           '5.8'],
        ['存储费（30天）',     '—'],
        ['收单费',             '15.0'],
        ['平台佣金（21%）',    '210.0'],
        ['广告费',             '42.9'],
        ['销售增值税（20%）',  '94.9'],
        ['杂费',               '2.0'],
        ['进项税抵扣',         '−71.5'],
        ['**总成本**',         '**881**'],
        ['售价（11681 RUB ÷ 11.7）', '998'],
        ['**利润**',           '**117（11.76%）**'],
    ],
    col_widths=[0.60, 0.40]
)

# ── 待确认问题 ────────────────────────────────────────────────────────────────
h2('七、待确认问题')

questions = [
    ('WB 运费算法',   'WB 运费 = 首升费用 + 第二升起费用，需确认各仓库的具体费率参数是否需要单独维护。'),
    ('汇率更新频率',  '汇率字段目前存在每一行，更新时需批量修改。如对实时性要求高，可考虑提取为独立配置表。'),
    ('历史成本留存',  '成本数据变化后是否需要保留历史版本？例如进货价涨价后，老库存应按旧成本还是新成本核算利润。'),
    ('OZON-Domos 账号', '该店铺目前未在数据库中，需补录凭证后方可在报表中填入真实数据。'),
]

for i, (title, desc) in enumerate(questions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(f'{i}. {title}：')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

# ── 保存 ──────────────────────────────────────────────────────────────────────
output_path = 'ozon_documents/SKU成本表结构设计.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
